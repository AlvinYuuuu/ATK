"""
Document parsing tools for extracting information from tender documents and other files.
"""

import os
import re
from typing import Dict, List, Any, Optional
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from a file (supports .txt, .md, .pdf, .docx).
    
    Args:
        file_path: Path to the file to extract text from
    
    Returns:
        Extracted text content
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        return f"Error: File {file_path} does not exist"
    
    # Handle different file types
    if file_path.suffix.lower() == '.txt' or file_path.suffix.lower() == '.md':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading file: {str(e)}"
    
    elif file_path.suffix.lower() == '.pdf':
        if not PDF_AVAILABLE:
            return f"PDF parsing not available. Please install PyMuPDF: pip install PyMuPDF"
        try:
            return extract_text_from_pdf(str(file_path))
        except Exception as e:
            return f"Error parsing PDF: {str(e)}"
    
    elif file_path.suffix.lower() == '.docx':
        if not DOCX_AVAILABLE:
            return f"DOCX parsing not available. Please install python-docx: pip install python-docx"
        try:
            return extract_text_from_docx(str(file_path))
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"
    
    else:
        return f"Unsupported file type: {file_path.suffix}"


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        Extracted text content
    """
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(f"Failed to parse DOCX file: {str(e)}")


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text content
    """
    try:
        doc = fitz.open(file_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_content.append(page.get_text())
        
        doc.close()
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(f"Failed to parse PDF file: {str(e)}")


def extract_requirements_from_text(text: str) -> Dict[str, Any]:
    """
    Extract requirements and key information from tender text.
    
    Args:
        text: The text content to analyze
    
    Returns:
        Dictionary containing extracted requirements and information
    """
    requirements = {
        "functional_requirements": [],
        "non_functional_requirements": [],
        "technical_constraints": [],
        "budget_constraints": [],
        "timeline_constraints": [],
        "client_background": "",
        "project_scope": "",
        "key_stakeholders": [],
        "success_criteria": []
    }
    
    # Extract functional requirements (basic pattern matching)
    functional_patterns = [
        r"must\s+(?:be\s+)?(?:able\s+to\s+)?(.+?)(?:\.|$)",
        r"should\s+(?:be\s+)?(?:able\s+to\s+)?(.+?)(?:\.|$)",
        r"shall\s+(?:be\s+)?(?:able\s+to\s+)?(.+?)(?:\.|$)",
        r"system\s+(?:must|should|shall)\s+(?:be\s+)?(?:able\s+to\s+)?(.+?)(?:\.|$)"
    ]
    
    for pattern in functional_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        requirements["functional_requirements"].extend(matches)
    
    # Extract budget information
    budget_patterns = [
        r"budget[:\s]+([^.\n]+)",
        r"cost[:\s]+([^.\n]+)",
        r"price[:\s]+([^.\n]+)",
        r"(\$[\d,]+)",
        r"([\d,]+\s*(?:USD|EUR|GBP))"
    ]
    
    for pattern in budget_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        requirements["budget_constraints"].extend(matches)
    
    # Extract timeline information
    timeline_patterns = [
        r"deadline[:\s]+([^.\n]+)",
        r"timeline[:\s]+([^.\n]+)",
        r"duration[:\s]+([^.\n]+)",
        r"(\d+\s*(?:weeks?|months?|days?))"
    ]
    
    for pattern in timeline_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        requirements["timeline_constraints"].extend(matches)
    
    # Extract technical constraints
    tech_patterns = [
        r"technology[:\s]+([^.\n]+)",
        r"platform[:\s]+([^.\n]+)",
        r"framework[:\s]+([^.\n]+)",
        r"language[:\s]+([^.\n]+)",
        r"database[:\s]+([^.\n]+)"
    ]
    
    for pattern in tech_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        requirements["technical_constraints"].extend(matches)
    
    # Extract client background (basic approach)
    client_sections = re.findall(r"(?:about|background|company|organization)[:\s]+([^.\n]+)", text, re.IGNORECASE)
    if client_sections:
        requirements["client_background"] = " ".join(client_sections)
    
    # Extract project scope
    scope_patterns = [
        r"scope[:\s]+([^.\n]+)",
        r"objective[:\s]+([^.\n]+)",
        r"goal[:\s]+([^.\n]+)"
    ]
    
    for pattern in scope_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            requirements["project_scope"] = " ".join(matches)
            break
    
    # Clean up lists (remove duplicates and empty items)
    for key in requirements:
        if isinstance(requirements[key], list):
            requirements[key] = list(set([item.strip() for item in requirements[key] if item.strip()]))
    
    return requirements


def identify_missing_information(text: str) -> List[str]:
    """
    Identify potentially missing information in the tender document.
    
    Args:
        text: The text content to analyze
    
    Returns:
        List of potentially missing information items
    """
    missing_info = []
    
    # Check for common missing elements
    checks = [
        ("budget", r"budget|cost|price|financial"),
        ("timeline", r"deadline|timeline|duration|schedule"),
        ("technical_specifications", r"technology|platform|framework|architecture"),
        ("success_criteria", r"success|criteria|metrics|kpi"),
        ("stakeholders", r"stakeholder|contact|person|team"),
        ("current_system", r"existing|current|legacy|system"),
        ("integration_requirements", r"integration|api|interface|connect"),
        ("security_requirements", r"security|authentication|authorization|encryption"),
        ("compliance_requirements", r"compliance|regulation|standard|certification")
    ]
    
    for item, pattern in checks:
        if not re.search(pattern, text, re.IGNORECASE):
            missing_info.append(f"Missing {item.replace('_', ' ')} information")
    
    return missing_info


def analyze_document_structure(text: str) -> Dict[str, Any]:
    """
    Analyze the structure and organization of the document.
    
    Args:
        text: The text content to analyze
    
    Returns:
        Dictionary containing structural analysis
    """
    analysis = {
        "sections": [],
        "word_count": len(text.split()),
        "has_executive_summary": False,
        "has_requirements_section": False,
        "has_technical_specifications": False,
        "has_timeline": False,
        "has_budget": False,
        "document_quality_score": 0
    }
    
    # Identify sections
    lines = text.split('\n')
    current_section = ""
    
    for line in lines:
        line = line.strip()
        if line and line.isupper() or re.match(r'^\d+\.\s+[A-Z]', line):
            current_section = line
            analysis["sections"].append(current_section)
    
    # Check for specific sections
    text_lower = text.lower()
    analysis["has_executive_summary"] = bool(re.search(r'executive\s+summary|overview|introduction', text_lower))
    analysis["has_requirements_section"] = bool(re.search(r'requirement|specification|scope', text_lower))
    analysis["has_technical_specifications"] = bool(re.search(r'technical|technology|architecture|system', text_lower))
    analysis["has_timeline"] = bool(re.search(r'timeline|schedule|deadline|duration', text_lower))
    analysis["has_budget"] = bool(re.search(r'budget|cost|price|financial', text_lower))
    
    # Calculate quality score
    score = 0
    if analysis["has_executive_summary"]: score += 20
    if analysis["has_requirements_section"]: score += 25
    if analysis["has_technical_specifications"]: score += 25
    if analysis["has_timeline"]: score += 15
    if analysis["has_budget"]: score += 15
    
    analysis["document_quality_score"] = score
    
    return analysis
